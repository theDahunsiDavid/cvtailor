import subprocess
import os
import re
import io
import json
import html
import pdfkit #note: ensure wkhtmltopdf is already installed
from django.conf import settings
from openai import OpenAI
client = OpenAI(api_key=settings.OPENAI_API_KEY)
from docx import Document
from docx.text.run import Run
from docx.text.paragraph import Paragraph
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from difflib import ndiff
from django.shortcuts import render, redirect
from django.contrib.auth import get_user_model, login
from django.contrib.auth import logout, authenticate
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse, HttpResponse
from django.urls import reverse
from django.contrib import messages
from .models import JobApplication
from .forms import JobApplicationForm

def home(request):
    return render(request, 'index.html')


def signUp(request):
    """
    Custom signup view for user registration.

    Handles:
        - GET requests to display the signupu form.
        - POST requests to process user registration.
    """
    User = get_user_model()

    if request.user.is_authenticated:
        return redirect(reverse('home'))

    next_url = request.GET.get('next', reverse('home'))

    if request.method == 'POST':
        email = request.POST.get('signup-email', '').strip()
        password = request.POST.get('signup-password', '').strip()
        confirm_password = request.POST.get(
            'confirm-password',
            ''
        ).strip()

        errors = {}

        if not email:
            errors['email'] = "Email is required."
        if not password:
            errors['password'] = "Password is required."
        if password != confirm_password:
            errors['confirm_password'] = "Passwords don't match."
        if User.objects.filter(email=email).exists():
            errors['email'] = "Email already used by someone else."

        if errors:
            return JsonResponse({
                'success': False,
                'errors': errors
            })

        user = User.objects.create_user(email=email, password=password)
        user.save()
        login(request, user)

        return JsonResponse({
            'success': True,
            'redirect_url': next_url
        })

    return JsonResponse({
        'success': False,
        'errors': {'general': "Invalid request method."}
    })


def signIn(request):
    """Logs the user into their account."""
    if request.method == 'POST':
        email = request.POST.get('login-email', '').strip()
        password = request.POST.get('login-password', '').strip()

        errors = {}

        if not email:
            errors['email'] = "Email is required."
        if not password:
            errors['password'] = "Password is required."

        if errors:
            return JsonResponse({
                'success': False,
                'errors': errors
            })

        user = authenticate(
            request,
            username=email,
            password=password
        )

        if user is not None:
            login(request, user)

            next_url = request.GET.get('next', reverse('home'))
            return JsonResponse({
                'success': True,
                'redirect_url': next_url
            })
        else:
            errors['general'] = "Invalid email or password."
            return JsonResponse({
                'success': False,
                'errors': errors
            })

    return JsonResponse({
        'success': False,
        'errors': {'general': "Invalid request method."}
    })


def signOut(request):
    """Logs the user out of their cvtailor account."""
    logout(request)
    return redirect('home')

def upload_cv(request):
    if request.method == 'POST':
        form = JobApplicationForm(request.POST, request.FILES)
        if form.is_valid():
            job_application = form.save()  # Save the job application to the database
            return process_uploaded_file(request, job_application.id)
    else:
        form = JobApplicationForm()

    return render(request, 'index.html', {'form': form})

def convert_docx_to_html(docx_file):
    """Converts a DOCX file to HTML."""
    try:
        if docx_file and docx_file.name.endswith('.docx'):
            doc = Document(docx_file)
            html_content = []
            
            doc_rels = {
                rid: rel.target_ref for rid, rel in doc.part.rels.items() 
                if rel.reltype == RT.HYPERLINK
            }
            
            in_list = False
            list_items = []
            
            for paragraph in doc.paragraphs:
                if not paragraph.text.strip() and not in_list:
                    html_content.append('<p><br></p>')
                    continue

                is_list_item = paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None
                
                if is_list_item:
                    if not in_list:
                        in_list = True
                        list_items = []
                    
                    item_content = []
                    
                    for element in paragraph._element:
                        if element.tag.endswith('hyperlink'):
                            rel_id = element.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                            if rel_id and rel_id in doc_rels:
                                url = doc_rels[rel_id]
                                texts = []
                                for run in element.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
                                    t_element = run.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                                    if t_element is not None and t_element.text:
                                        texts.append(html.escape(t_element.text))
                                
                                hyperlink_text = ''.join(texts)
                                if hyperlink_text:  # Only add if we have text
                                    item_content.append(f'<a href="{url}">{hyperlink_text}</a>')
                        elif element.tag.endswith('r'):
                            run_text = ''.join(
                                html.escape(child.text or '') 
                                for child in element.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                            )
                            rPr = element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
                            if rPr is not None:
                                if rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}b') is not None:
                                    run_text = f'<strong>{run_text}</strong>'
                                if rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}i') is not None:
                                    run_text = f'<em>{run_text}</em>'
                            item_content.append(run_text)
                    
                    list_items.append(''.join(item_content))
                
                elif in_list:
                    html_content.append('<ul>\n' + '\n'.join([f'<li>{item}</li>' for item in list_items]) + '\n</ul>')
                    list_items = []
                    in_list = False
                    
                    para_html = ['<p>']
                    
                    for element in paragraph._element:
                        if element.tag.endswith('hyperlink'):
                            rel_id = element.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                            if rel_id and rel_id in doc_rels:
                                url = doc_rels[rel_id]
                                texts = []
                                for run in element.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
                                    t_element = run.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                                    if t_element is not None and t_element.text:
                                        texts.append(html.escape(t_element.text))
                                
                                hyperlink_text = ''.join(texts)
                                if hyperlink_text:  # Only add if we have text
                                    para_html.append(f'<a href="{url}">{hyperlink_text}</a>')
                        elif element.tag.endswith('r'):
                            run_text = ''.join(
                                html.escape(child.text or '') 
                                for child in element.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                            )
                            rPr = element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
                            if rPr is not None:
                                if rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}b') is not None:
                                    run_text = f'<strong>{run_text}</strong>'
                                if rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}i') is not None:
                                    run_text = f'<em>{run_text}</em>'
                            para_html.append(run_text)
                    
                    para_html.append('</p>')
                    html_content.append(''.join(para_html))
                
                else:
                    para_html = ['<p>']
                    
                    for element in paragraph._element:
                        if element.tag.endswith('hyperlink'):
                            rel_id = element.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                            if rel_id and rel_id in doc_rels:
                                url = doc_rels[rel_id]
                                texts = []
                                for run in element.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
                                    t_element = run.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                                    if t_element is not None and t_element.text:
                                        texts.append(html.escape(t_element.text))
                                
                                hyperlink_text = ''.join(texts)
                                if hyperlink_text:  # Only add if we have text
                                    para_html.append(f'<a href="{url}">{hyperlink_text}</a>')
                        elif element.tag.endswith('r'):
                            run_text = ''.join(
                                html.escape(child.text or '') 
                                for child in element.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                            )
                            rPr = element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
                            if rPr is not None:
                                if rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}b') is not None:
                                    run_text = f'<strong>{run_text}</strong>'
                                if rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}i') is not None:
                                    run_text = f'<em>{run_text}</em>'
                            para_html.append(run_text)
                    
                    para_html.append('</p>')
                    html_content.append(''.join(para_html))

            if in_list and list_items:
                html_content.append('<ul>\n' + '\n'.join([f'<li>{item}</li>' for item in list_items]) + '\n</ul>')
            
            resume_text = '\n'.join(html_content)
            return resume_text

        
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        return None

def process_uploaded_file(request, file_id):
    """Processes the uploaded CV and converts to HTML."""
    job_application = JobApplication.objects.get(id=file_id)
    file_path = job_application.uploaded_file.path
    job_description = job_application.job_description

    try:
        with open(file_path, 'rb') as file:
            html_content = convert_docx_to_html(file)
        if html_content is None:
            return HttpResponse("Error converting the document.", status=500)

        sections = {"CV Content": html_content}

        return render(request, 'preview.html', {
            'sections': sections,
            'job_description': job_description
        })

    except FileNotFoundError as e:
        return HttpResponse(f"File not found: {e}", status=404)
    except Exception as e:
        return HttpResponse(f"An error occurred: {e}", status=500)

def convert_to_ats_format(request):
    if request.method == "POST":
        try:
            cv_content = request.POST.get("cv_content")
            job_description = request.POST.get("job_description")

            ats_cv_content, match_score, suggestions = generate_ats_format_and_match_score(cv_content, job_description)

            return JsonResponse({
                "ats_cv_content": ats_cv_content,
                "match_score": match_score,
                "suggestions": suggestions
                })

        except Exception as e:
            return JsonResponse({'error': f"An error occurred: {str(e)}"}, status=500)

def generate_ats_format_and_match_score(cv_content, job_description):
    try:
        prompt = (
            f"First, proofread the following CV content, then convert it to an ATS-friendly format as possible without removing personal data like name, address, email, contact info, and other identifying details that would normally not be in an ATS-friendly CV. Let there be no omissions or missing parts, meaning do not remove any sections containing personal information or contact information. Then proofread and edit the content to remove all typographical errors."
            f"Section titles/headings in the CV content might be capitalized. If the section titles/headings are not in capital letters, then capitalize them. Separate all sections with an empty line only, not with commas or separators. All sections in CV content must be present in the ATS-friendly format, including all responsibilities under each work experience entry. Ensure no sections or responsibilities are missing in the ATS-friendly format."
            f"Do NOT remove a section because its title/heading is a non-standard title. However, where you can, rename the non-standard titles (e.g., PROFILE SUMMARY) for each section to standard, clear section headings like PROFESSIONAL SUMMARY, WORK EXPERIENCE, SKILLS, EDUCATION, etc. Capitalize all section titles/headings."
            f"Remove any empty lines between a section title/heading and its content."
            f"Do not separate sections with '---' or the likes."
            f"Do not end your response with '---' or the likes."
            f"If the first line is a name, capitalize it, and put other personal data like name, address, email, contact info, and other identifying details under the capitalized name with no empty line between them."
            f"Format each job entry as 'Company Name (Position) [Date Range]' and include a bullet point list of responsibilities under each position. Do not format each job entry with a bullet point like '- Company Name (Position) [Date Range]'. Use hyphens in the Date Ranges. Do not use en dash or em dash in the Date Ranges."
            f"Format the WORK EXPERIENCE section in reverse-chronological order."
            f"Use simple headings for section titles, and ensure all sections are clearly separated without unnecessary commas or separators."
            f"Format the WORK EXPERIENCE section in reverse-chronological format."
            f"Always include a Match Score (out of 100) that indicates how closely the CV matches the job description as exactly 'Match Score: X/100' where X is a number and is the Match Score. Always include numbered suggestions (changes that can be made to the ATS-friendly CV to improve the Match Score) as 'Suggestions to improve Match Score: Y' where Y is the numbered suggestions."
            f"Proofread and edit the contents of all the sections. Fix all the grammatical errors or typos in the contents of all the sections."
            f"Capitalize all section titles. That is, capitalize all section headers."
            f"Double-check that all sections and responsibilities from the original CV are included in your response.\n"
            f"CV Content: {cv_content}\n"
            f"Job Description: {job_description}\n"
)

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "user", "content": prompt}
            ],
        )

        ats_content = response.choices[0].message.content

        suggestions = extract_suggestions(ats_content)
        match_score_match = re.search(r'Match Score:\s*(\d+)\s*/\s*100', ats_content, re.IGNORECASE)
        if match_score_match:
            match_score = match_score_match.group(1)
            ats_content = ats_content[:match_score_match.start()].strip()
        else:
            match_score = "N/A"

        ats_content = apply_custom_formatting(ats_content)

        formatted_ats_content = format_for_quill(ats_content)
        return formatted_ats_content, match_score, suggestions

    except Exception as e:
        print(f"Error generating ATS content and match score: {e}")
        return "Error generating ATS content", "N/A"

def format_for_quill(content):
    """
    Format text to be displayed properly in the Quill editor.
    Preserves paragraphs, line breaks, and lists.
    """
    lines = content.split('\n')
    html_lines = []
    in_list = False
    list_type = None
    list_items = []

    for line in lines:
        line = line.strip()
        if not line:
            if in_list:
                html_lines.append(f'<{list_type}>\n' + '\n'.join([f'<li>{item}</li>' for item in list_items]) + f'\n</{list_type}>')
                list_items = []
                in_list = False
                list_type = None
            html_lines.append('<p><br></p>')
            continue
        
        if line.startswith('-') or line.startswith('•') or line.startswith('*'):
            if not in_list:
                in_list = True
                list_type = 'ul'
            item_text = line[1:].strip()
            list_items.append(item_text)
        elif re.match(r'^\d+[\.\)]', line):
            if not in_list:
                in_list = True
                list_type = 'ol'  # Set list type to ordered
            item_text = re.sub(r'^\d+[\.\)]\s*', '', line)
            list_items.append(item_text)
        else:
            if in_list:
                html_lines.append(f'<{list_type}>\n' + '\n'.join([f'<li>{item}</li>' for item in list_items]) + f'\n</{list_type}>')
                list_items = []
                in_list = False
                list_type = None
            
            if line.isupper() or (line[0].isupper() and ':' in line):
                html_lines.append(f'<h3>{line}</h3>')
            else:
                html_lines.append(f'<p>{line}</p>')
    
    if in_list:
        html_lines.append(f'<{list_type}>\n' + '\n'.join([f'<li>{item}</li>' for item in list_items]) + f'\n</{list_type}>')
    
    return '\n'.join(html_lines)


def apply_custom_formatting(content):
    """
    Format the content to remove asterisks, convert them to headings,
    and replace bullet points with dashes.
    """

    # Converts lines like **PROFILE SUMMARY** into PROFILE SUMMARY
    content = re.sub(r'\*\*(.*?)\*\*', r'\n\1\n', content)

    # Replaces bullet points (asterisks or other symbols) with dashes
    content = content.replace('•', '-').replace('*', '-')

    # Removes empty line at the top of the text, if it exists
    content = re.sub(r'^\s*\n', '', content)

    # Removes excess newlines but keeps paragraph spacing
    content = re.sub(r'\n\s*\n', '\n\n', content)

    return content

def extract_suggestions(ats_response):
    # Example: Assuming the response contains a section with suggestions
    suggestions_section = re.search(r'Suggestions to improve Match Score:\s*(.*?)(?=\n\n|\Z)', ats_response, re.DOTALL | re.IGNORECASE)
    if suggestions_section:
        suggestions_text = suggestions_section.group(1).strip()
        return [suggestion.strip() for suggestion in suggestions_text.split("\n")]  # Split suggestions into a list
    return []

"""
@csrf_exempt  # Add CSRF exemption if needed
def apply_suggestion(request):
    # View to apply a suggestion to the current CV content.
    if request.method == "POST":
        try:
            # Extract current CV content and the suggestion from the request
            cv_content = request.POST.get("cv_content")
            suggestion = request.POST.get("suggestion")

            # Apply the suggestion to the CV content
            updated_cv_content = apply_suggestion_to_cv(cv_content, suggestion)

            return JsonResponse({"updated_cv_content": updated_cv_content})

        except Exception as e:
            return JsonResponse({"error": f"An error occurred: {str(e)}"}, status=500)

def apply_suggestion_to_cv(cv_content, suggestion):

    Applies a specific suggestion to the CV content.

    Args:
        cv_content (str): The current CV content.
        suggestion (str): The suggestion to be applied.

    Returns:
        str: The updated CV content after applying the suggestion.

    try:
        # Here, implement the logic to modify the CV based on the suggestion.
        # You might want to use different rules based on the suggestion text.
        # For example, replacing certain phrases, reformatting sections, etc.

        # Placeholder logic for now:
        if "add experience" in suggestion.lower():
            # Example: Adding an experience section suggestion
            cv_content += "<br><br><strong>Additional Experience:</strong> Suggested addition goes here."
        elif "remove redundancy" in suggestion.lower():
            # Example: Removing redundancy as a suggestion
            cv_content = cv_content.replace("redundant phrase", "")

        # Add more conditions based on the type of suggestions you receive
        # ...

        return cv_content

    except Exception as e:
        print(f"Error applying suggestion: {e}")
        return cv_content  # Return unmodified content if there's an error

def download_docx(request):
    # Get the tailored CV content from the request
    cv_content = request.POST.get('cv_content', '')

    # Create a DOCX document
    doc = Document()
    doc.add_paragraph(cv_content)

    # Prepare the DOCX file to be downloaded
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Create an HTTP response with the DOCX file
    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="tailored_cv.docx"'

    return response
"""

def download_pdf(request):
    cv_content = request.POST.get('cv_content', '')
    
    html_document = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            body {{ font-family: Arial, sans-serif; margin: 20px; }}
            ul {{ list-style-type: disc; }}
            ol {{ list-style-type: decimal; }}
            li {{ margin-bottom: 5px; }}
        </style>
    </head>
    <body>
        {cv_content}
    </body>
    </html>
    """

    pdf = pdfkit.from_string(html_document, False, options={
        'encoding': "UTF-8",
        'quiet': '',
        'margin-top': '20mm',
        'margin-right': '20mm',
        'margin-bottom': '20mm',
        'margin-left': '20mm',
    })

    response = HttpResponse(pdf, content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="tailored_cv.pdf"'

    return response

@csrf_exempt
def implement_suggestion(request):
    """Processes the request to implement suggestion in CV text."""
    data = json.loads(request.body)
    text = data['text']
    suggestion = data['suggestion']

    prompt = (
        f"Please modify the provided CV text based on the user's suggestion without altering unrelated sections."
        f"Just give me all the CV text (including all the modifications) ONLY in your response."
    )

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": prompt},
            {"role": "user", "content": f"CV Text: {text}\n\nSuggestion: {suggestion}"}
        ]
    )

    modified_text = response.choices[0].message.content

    highlighted_text = highlight_changes(text, modified_text)

    print("Highlighted Text:", highlighted_text)

    return JsonResponse({"highlighted_text": highlighted_text})

def highlight_changes(original_text: str, modified_text: str) -> str:
    """
    Compares original and modified text & highlights differences-
    -with Accept and Reject controls.

    Args:
        original_text (str): original CV text.
        modified_text (str): modified CV text w/ suggestion implemented.

    Returns:
        str: modified text w/ only changes highlighted & controls added.
    """
    highlighted_text = ""
    diff = ndiff(original_text.split(), modified_text.split())
    change_id = 0

    for token in diff:
        if token.startswith("+ "):
            highlighted_text += (
                f'<span class="highlighted-change" id="change-{change_id}">'
                f'{token[2:]}'
                f'<button onclick="acceptChange(\'change-{change_id}\')">Accept</button>'
                f'<button onclick="rejectChange(\'change-{change_id}\')">Reject</button>'
                f'</span> '
            )
            change_id += 1
        elif token.startswith("- "):
            continue
        else:
            highlighted_text += token[2:] + " "

    return highlighted_text.strip()

